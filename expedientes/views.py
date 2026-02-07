import os
import json
import uuid
import zipfile
from io import BytesIO
from datetime import timedelta
from decimal import Decimal 
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth import logout
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.db.models import Count, Q
from django.http import JsonResponse, HttpResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.files.base import ContentFile
from django.utils import timezone
from django.utils.text import slugify 
from django.template.loader import render_to_string
from django.core.mail import EmailMultiAlternatives
from django.conf import settings 
from django.utils.html import strip_tags
from email.mime.image import MIMEImage
# Importante para serializar los servicios en la nueva cotizaci√≥n
from django.core.serializers import serialize 
from .models import Cliente
from django.shortcuts import get_object_or_404, redirect
from django.contrib import messages
from .models import Carpeta, Documento, Cliente
# Librer√≠as para Documentos
from docxtpl import DocxTemplate
import mammoth
from docx import Document as DocumentoWord 
import weasyprint 
from django.core.mail import send_mail

import qrcode
from io import BytesIO
import base64
from qrcode.image.styledpil import StyledPilImage
from qrcode.image.styles.moduledrawers import RoundedModuleDrawer

# Importaci√≥n de Modelos
from .models import (
    Usuario, Cliente, Carpeta, Expediente, Documento, 
    Tarea, Bitacora, Plantilla, VariableEstandar,
    Servicio, Cotizacion, ItemCotizacion, PlantillaMensaje,
    CuentaPorCobrar, Pago, Evento, CampoAdicional,Archivo,
)

from decimal import Decimal

# ==========================================
# 1. AUTENTICACI√ìN Y PERFIL
# ==========================================

def signout(request):
    logout(request)
    return redirect('login')

def registro(request):
    if request.method == 'POST':
        first_name = request.POST.get('first_name')
        last_name = request.POST.get('last_name')
        username = request.POST.get('username')
        email = request.POST.get('email')
        pass1 = request.POST.get('pass1')
        pass2 = request.POST.get('pass2')

        if pass1 != pass2:
            messages.error(request, "Las contrase√±as no coinciden.")
            return render(request, 'registro.html')

        if Usuario.objects.filter(username=username).exists():
            messages.error(request, "El usuario ya existe.")
            return render(request, 'registro.html')

        try:
            Usuario.objects.create_user(
                username=username, email=email, password=pass1,
                first_name=first_name, last_name=last_name, is_active=False
            )
            return render(request, 'registro_pendiente.html')
        except Exception as e:
            messages.error(request, f"Error del sistema: {e}")

    return render(request, 'registro.html')

@login_required
def mi_perfil(request):
    user = request.user
    if request.method == 'POST':
        user.first_name = request.POST.get('first_name')
        user.last_name = request.POST.get('last_name')
        user.email = request.POST.get('email')
        user.telefono = request.POST.get('telefono')
        user.puesto = request.POST.get('puesto')
        
        if request.FILES.get('avatar'):
            user.avatar = request.FILES['avatar']
            
        user.save()
        messages.success(request, "Perfil actualizado correctamente.")
        return redirect('mi_perfil')
    return render(request, 'usuarios/mi_perfil.html', {'user': user})

# ==========================================
# 2. GESTI√ìN DE USUARIOS (ADMIN)
# ==========================================

@login_required
def gestion_usuarios(request):
    if request.user.rol != 'admin': return redirect('dashboard')
    usuarios = Usuario.objects.all().order_by('-date_joined')
    return render(request, 'gestion_usuarios.html', {'usuarios': usuarios})

@login_required
def autorizar_usuario(request, user_id):
    if request.user.rol != 'admin': return redirect('dashboard')
    user = get_object_or_404(Usuario, id=user_id)
    user.is_active = True
    user.save()
    messages.success(request, f"Usuario {user.username} autorizado.")
    return redirect('gestion_usuarios')

@login_required
def editar_usuario(request, user_id):
    if request.user.rol != 'admin': return redirect('dashboard')
    user_obj = get_object_or_404(Usuario, id=user_id)
    clientes_disponibles = Cliente.objects.all().order_by('nombre_empresa')
    
    if request.method == 'POST':
        user_obj.rol = request.POST.get('rol')
        user_obj.first_name = request.POST.get('first_name') or ""
        user_obj.last_name = request.POST.get('last_name') or ""
        user_obj.email = request.POST.get('email')
        user_obj.telefono = request.POST.get('telefono') or None
        user_obj.puesto = request.POST.get('puesto') or None
        
        # Permisos booleanos
        user_obj.can_create_client = request.POST.get('can_create_client') == 'on'
        user_obj.can_edit_client = request.POST.get('can_edit_client') == 'on'
        user_obj.can_delete_client = request.POST.get('can_delete_client') == 'on'
        user_obj.can_upload_files = request.POST.get('can_upload_files') == 'on'
        user_obj.can_view_documents = request.POST.get('can_view_documents') == 'on'
        user_obj.can_manage_users = request.POST.get('can_manage_users') == 'on'

        # Accesos a m√≥dulos
        user_obj.access_finanzas = request.POST.get('access_finanzas') == 'on'
        user_obj.access_cotizaciones = request.POST.get('access_cotizaciones') == 'on'
        user_obj.access_contratos = request.POST.get('access_contratos') == 'on'
        user_obj.access_disenador = request.POST.get('access_disenador') == 'on'
        user_obj.access_agenda = request.POST.get('access_agenda') == 'on'
        
        clientes_ids = request.POST.getlist('clientes_asignados')
        user_obj.save()
        
        if user_obj.rol != 'admin':
            user_obj.clientes_asignados.set(clientes_ids)
        else:
            user_obj.clientes_asignados.clear()
            
        messages.success(request, f"Permisos de {user_obj.username} actualizados.")
        return redirect('gestion_usuarios')

    return render(request, 'usuarios/editar_usuario.html', {'u': user_obj, 'clientes': clientes_disponibles})

@login_required
def eliminar_usuario(request, user_id):
    if request.user.rol != 'admin': return redirect('dashboard')
    u = get_object_or_404(Usuario, id=user_id)
    if u == request.user:
        messages.error(request, "No puedes eliminarte a ti mismo.")
        return redirect('gestion_usuarios')
    u.delete()
    messages.success(request, "Usuario eliminado.")
    return redirect('gestion_usuarios')

# ==========================================
# 3. DASHBOARD Y CLIENTES
# ==========================================

@login_required
def dashboard(request):
    if request.user.rol == 'admin':
        mis_clientes = Cliente.objects.all()
    else:
        mis_clientes = request.user.clientes_asignados.all()

    stats = {
        'total_clientes': mis_clientes.count(),
        'expedientes_activos': Expediente.objects.filter(cliente__in=mis_clientes, estado='abierto').count(),
        'tareas_pendientes': Tarea.objects.filter(cliente__in=mis_clientes, completada=False).count(),
        'docs_subidos': Documento.objects.filter(cliente__in=mis_clientes).count()
    }
    
    hoy = timezone.now().date()
    tareas_criticas = Tarea.objects.filter(cliente__in=mis_clientes, completada=False, fecha_limite__lte=hoy)
    
    clientes = mis_clientes.annotate(
        num_expedientes=Count('expedientes', distinct=True),
        urgencias=Count('tareas', filter=Q(tareas__prioridad='alta', tareas__completada=False), distinct=True)
    ).order_by('-urgencias', '-fecha_registro')

    pendientes = 0
    if request.user.rol == 'admin':
        pendientes = Usuario.objects.filter(is_active=False).count()

    return render(request, 'dashboard.html', {
        'clientes': clientes,
        'stats': stats,
        'usuarios_pendientes_conteo': pendientes,
        'now': timezone.now(),
        'alertas': {'tareas': tareas_criticas} 
    })

@login_required
def nuevo_cliente(request):
    if not (request.user.rol == 'admin' or request.user.can_create_client):
        return redirect('dashboard')
    if request.method == 'POST':
        c = Cliente.objects.create(
            nombre_empresa=request.POST.get('nombre_empresa'),
            nombre_contacto=request.POST.get('nombre_contacto'),
            email=request.POST.get('email'),
            telefono=request.POST.get('telefono'),
            logo=request.FILES.get('logo')
        )
        if request.user.rol != 'admin':
            request.user.clientes_asignados.add(c)
        return redirect('dashboard')
    return render(request, 'nuevo_cliente.html')

@login_required
def eliminar_cliente(request, cliente_id):
    if request.user.rol != 'admin' and not request.user.can_delete_client:
        return redirect('dashboard')
    cliente = get_object_or_404(Cliente, id=cliente_id)
    cliente.delete()
    messages.success(request, "Cliente eliminado.")
    return redirect('dashboard')

@login_required
def detalle_cliente(request, cliente_id, carpeta_id=None):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    
    if request.user.rol != 'admin' and cliente not in request.user.clientes_asignados.all():
        messages.error(request, "‚õî Acceso Denegado.")
        return redirect('dashboard')

    carpeta_actual = None
    breadcrumbs = []
    
    if carpeta_id:
        carpeta_actual = get_object_or_404(Carpeta, id=carpeta_id, cliente=cliente)
        crumb = carpeta_actual
        while crumb:
            breadcrumbs.insert(0, crumb)
            crumb = crumb.padre

    if carpeta_actual:
        carpetas = cliente.carpetas_drive.filter(padre=carpeta_actual)
        documentos = cliente.documentos_cliente.filter(carpeta=carpeta_actual)
    else:
        carpetas = cliente.carpetas_drive.filter(padre__isnull=True)
        documentos = cliente.documentos_cliente.filter(carpeta__isnull=True)

    stats_cliente = {
        'total_docs': cliente.documentos_cliente.count(),
        'expedientes_activos': cliente.expedientes.filter(estado='abierto').count(),
    }
    
    historial = Bitacora.objects.filter(cliente=cliente).select_related('usuario').order_by('-fecha')
    todas_carpetas = cliente.carpetas_drive.all()

    return render(request, 'detalle_cliente.html', {
        'cliente': cliente,
        'carpeta_actual': carpeta_actual,
        'breadcrumbs': breadcrumbs,
        'carpetas': carpetas,
        'documentos': documentos,
        'stats_cliente': stats_cliente,
        'historial': historial,
        'todas_carpetas': todas_carpetas,
    })

@login_required
def editar_cliente(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    if request.user.rol != 'admin' and cliente not in request.user.clientes_asignados.all():
        return redirect('dashboard')

    campos_dinamicos = CampoAdicional.objects.all()

    if request.method == 'POST':
        cliente.nombre_empresa = request.POST.get('nombre_empresa')
        cliente.nombre_contacto = request.POST.get('nombre_contacto')
        cliente.email = request.POST.get('email')
        cliente.telefono = request.POST.get('telefono')
        
        if request.FILES.get('logo'):
            cliente.logo = request.FILES['logo']

        datos_nuevos = cliente.datos_extra or {}
        for campo in campos_dinamicos:
            valor = request.POST.get(f"custom_{campo.id}")
            if valor:
                datos_nuevos[campo.nombre] = valor
        
        cliente.datos_extra = datos_nuevos
        cliente.save()
        Bitacora.objects.create(usuario=request.user, cliente=cliente, accion='edicion', descripcion="Actualiz√≥ datos.")
        messages.success(request, "Cliente actualizado.")
        return redirect('detalle_cliente', cliente_id=cliente.id)

    return render(request, 'clientes/editar.html', {
        'c': cliente,
        'campos_dinamicos': campos_dinamicos,
        'datos_existentes': cliente.datos_extra
    })

# ==========================================
# 4. CONFIGURACI√ìN Y DRIVE
# ==========================================

@login_required
def configurar_campos(request):
    if request.user.rol != 'admin': return redirect('dashboard')
    campos = CampoAdicional.objects.all()
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        if not CampoAdicional.objects.filter(nombre__iexact=nombre).exists():
            CampoAdicional.objects.create(nombre=nombre, tipo=request.POST.get('tipo'))
            messages.success(request, f"Campo '{nombre}' agregado.")
        return redirect('configurar_campos')
    return render(request, 'clientes/configurar_campos.html', {'campos': campos})

@login_required
def eliminar_campo_dinamico(request, campo_id):
    if request.user.rol != 'admin': return redirect('dashboard')
    get_object_or_404(CampoAdicional, id=campo_id).delete()
    return redirect('configurar_campos')

@login_required
def crear_carpeta(request, cliente_id):
    if request.method == 'POST':
        padre_id = request.POST.get('padre_id')
        padre = get_object_or_404(Carpeta, id=padre_id) if padre_id else None
        Carpeta.objects.create(nombre=request.POST.get('nombre'), cliente_id=cliente_id, padre=padre)
        if padre: return redirect('detalle_carpeta', cliente_id=cliente_id, carpeta_id=padre.id)
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
def eliminar_carpeta(request, carpeta_id):
    if not (request.user.can_delete_client or request.user.rol == 'admin'): return redirect('dashboard')
    c = get_object_or_404(Carpeta, id=carpeta_id)
    url_destino = 'detalle_carpeta' if c.padre else 'detalle_cliente'
    kwargs = {'cliente_id': c.cliente.id}
    if c.padre: kwargs['carpeta_id'] = c.padre.id
    c.delete()
    return redirect(url_destino, **kwargs)

@login_required
def crear_expediente(request, cliente_id):
    if request.method == 'POST':
        f = Carpeta.objects.create(nombre=f"EXP {request.POST.get('num_expediente')}: {request.POST.get('titulo')}", cliente_id=cliente_id, es_expediente=True)
        Expediente.objects.create(cliente_id=cliente_id, num_expediente=request.POST.get('num_expediente'), titulo=request.POST.get('titulo'), carpeta=f)
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
def subir_archivo_drive(request, cliente_id):
    if not (request.user.can_upload_files or request.user.rol == 'admin'): return redirect('detalle_cliente', cliente_id=cliente_id)
    if request.method == 'POST':
        cliente = get_object_or_404(Cliente, id=cliente_id)
        archivos = request.FILES.getlist('archivo')
        carpeta_raiz_id = request.POST.get('carpeta_id')
        carpeta_raiz = get_object_or_404(Carpeta, id=carpeta_raiz_id) if carpeta_raiz_id else None
        
        count = 0
        for f in archivos:
            Documento.objects.create(cliente_id=cliente_id, archivo=f, nombre_archivo=f.name, carpeta=carpeta_raiz, subido_por=request.user)
            count += 1
        
        ubicacion = carpeta_raiz.nombre if carpeta_raiz else "Ra√≠z"
        Bitacora.objects.create(usuario=request.user, cliente=cliente, accion='subida', descripcion=f"Subi√≥ {count} archivos en '{ubicacion}'.")
        if carpeta_raiz: return redirect('detalle_carpeta', cliente_id=cliente_id, carpeta_id=carpeta_raiz.id)
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
def eliminar_archivo_drive(request, archivo_id):
    doc = get_object_or_404(Documento, id=archivo_id)
    if not (request.user.can_delete_client or request.user.rol == 'admin'): return redirect('detalle_cliente', cliente_id=doc.cliente.id)
    c_id, padre_id = doc.cliente.id, doc.carpeta.id if doc.carpeta else None
    Bitacora.objects.create(usuario=request.user, cliente=doc.cliente, accion='eliminacion', descripcion=f"Elimin√≥ {doc.nombre_archivo}")
    doc.archivo.delete(); doc.delete()
    if padre_id: return redirect('detalle_carpeta', cliente_id=c_id, carpeta_id=padre_id)
    return redirect('detalle_cliente', cliente_id=c_id)

@login_required
def descargar_carpeta_zip(request, carpeta_id):
    carpeta = get_object_or_404(Carpeta, id=carpeta_id)
    if request.user.rol != 'admin' and carpeta.cliente not in request.user.clientes_asignados.all():
        return HttpResponse("Acceso Denegado", status=403)
    
    buffer = BytesIO()
    with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
        for file in Documento.objects.filter(carpeta=carpeta):
            try: zip_file.writestr(file.nombre_archivo, file.archivo.read())
            except: pass
    
    Bitacora.objects.create(usuario=request.user, cliente=carpeta.cliente, accion='descarga', descripcion=f"Descarg√≥ ZIP: {carpeta.nombre}")
    buffer.seek(0)
    response = HttpResponse(buffer, content_type='application/zip')
    response['Content-Disposition'] = f'attachment; filename="{carpeta.nombre}.zip"'
    return response

@login_required
def acciones_masivas_drive(request):
    if request.method == 'POST':
        accion = request.POST.get('accion')
        doc_ids = request.POST.getlist('doc_ids')
        docs = Documento.objects.filter(id__in=doc_ids)
        if not docs: return redirect(request.META.get('HTTP_REFERER'))
        
        cliente = docs.first().cliente
        if accion == 'eliminar':
            if not (request.user.can_delete_client or request.user.rol == 'admin'): return redirect(request.META.get('HTTP_REFERER'))
            count = docs.count()
            for doc in docs: doc.archivo.delete(); doc.delete()
            Bitacora.objects.create(usuario=request.user, cliente=cliente, accion='eliminacion', descripcion=f"Elimin√≥ {count} archivos masivamente.")
            messages.success(request, f"Se eliminaron {count} archivos.")
        
        elif accion == 'descargar':
            buffer = BytesIO()
            with zipfile.ZipFile(buffer, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                for doc in docs:
                    try: zip_file.writestr(doc.nombre_archivo, doc.archivo.read())
                    except: pass
            Bitacora.objects.create(usuario=request.user, cliente=cliente, accion='descarga', descripcion=f"Descarg√≥ selecci√≥n ZIP.")
            buffer.seek(0)
            response = HttpResponse(buffer, content_type='application/zip')
            response['Content-Disposition'] = f'attachment; filename="Seleccion.zip"'
            return response
            
    return redirect(request.META.get('HTTP_REFERER'))

@login_required
def preview_archivo(request, documento_id):
    doc = get_object_or_404(Documento, id=documento_id)
    ext = doc.nombre_archivo.split('.')[-1].lower()
    data = {'tipo': 'unknown', 'url': doc.archivo.url, 'nombre': doc.nombre_archivo}
    if ext in ['jpg', 'jpeg', 'png', 'gif', 'webp']: data['tipo'] = 'imagen'
    elif ext == 'pdf': data['tipo'] = 'pdf'
    elif ext == 'docx':
        data['tipo'] = 'docx'
        try:
            with doc.archivo.open() as f: data['html'] = mammoth.convert_to_html(f).value
        except: data['html'] = "Error de lectura."
    return JsonResponse(data)

# ==========================================
# 5. TAREAS
# ==========================================

@login_required
def gestionar_tarea(request, cliente_id):
    if request.method == 'POST':
        Tarea.objects.create(
            cliente_id=cliente_id, titulo=request.POST.get('titulo'),
            fecha_limite=request.POST.get('fecha_limite'), prioridad=request.POST.get('prioridad')
        )
    return redirect('detalle_cliente', cliente_id=cliente_id)

@login_required
def toggle_tarea(request, tarea_id):
    t = get_object_or_404(Tarea, id=tarea_id)
    t.completada = not t.completada
    t.save()
    return redirect('detalle_cliente', cliente_id=t.cliente.id)

@login_required
def editar_tarea(request, tarea_id):
    t = get_object_or_404(Tarea, id=tarea_id)
    if request.method == 'POST':
        t.titulo = request.POST.get('titulo')
        t.fecha_limite = request.POST.get('fecha_limite')
        t.prioridad = request.POST.get('prioridad')
        t.save()
    return redirect('detalle_cliente', cliente_id=t.cliente.id)

@login_required
def eliminar_tarea(request, tarea_id):
    t = get_object_or_404(Tarea, id=tarea_id)
    c_id = t.cliente.id
    t.delete()
    return redirect('detalle_cliente', cliente_id=c_id)

# ==========================================
# 6. CONTRATOS Y DISE√ëADOR
# ==========================================

@login_required
def generador_contratos(request, cliente_id):
    if not request.user.access_contratos: return redirect('dashboard')
    cliente = get_object_or_404(Cliente, id=cliente_id)
    
    if request.method == 'GET' and 'plantilla_id' not in request.GET:
        return render(request, 'generador/seleccionar.html', {
            'cliente': cliente,
            'plantillas': Plantilla.objects.all().order_by('-fecha_subida'),
            'glosario': VariableEstandar.objects.all().order_by('clave')
        })

    plantilla = get_object_or_404(Plantilla, id=request.GET.get('plantilla_id') or request.POST.get('plantilla_id'))
    doc = DocxTemplate(plantilla.archivo.path)
    vars_en_doc = doc.get_undeclared_template_variables()
    memoria = cliente.datos_extra if isinstance(cliente.datos_extra, dict) else {}
    formulario = []
    
    mapeo = {
        'cliente.nombre_empresa': cliente.nombre_empresa,
        'cliente.nombre_contacto': cliente.nombre_contacto,
        'cliente.email': cliente.email,
        'cliente.telefono': cliente.telefono,
        'fecha_actual': timezone.now().strftime("%d/%m/%Y"),
    }

    for v in vars_en_doc:
        var_std = VariableEstandar.objects.filter(clave=v).first()
        val = ""
        auto = False
        desc = "Variable"
        tipo = "text"

        if var_std:
            desc = var_std.descripcion
            if var_std.tipo == 'fecha': tipo = 'date'
            if var_std.origen == 'sistema':
                val = mapeo.get(var_std.campo_bd, '')
                auto = True
            else: val = memoria.get(v, '')
        else: val = memoria.get(v, '')

        formulario.append({'clave': v, 'valor': val, 'descripcion': desc, 'es_automatico': auto, 'tipo': tipo})

    if request.method == 'POST':
        contexto = {}
        nuevos_datos = {}
        for item in formulario:
            if item['es_automatico']: val = item['valor']
            else:
                val = request.POST.get(item['clave'], '').strip()
                nuevos_datos[item['clave']] = val
            contexto[item['clave']] = val
            
        cliente.datos_extra.update(nuevos_datos)
        cliente.save(update_fields=['datos_extra'])
        
        doc.render(contexto)
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        nombre = request.POST.get('nombre_archivo_salida', '').strip() or f"{plantilla.nombre} - {cliente.nombre_empresa}"
        if not nombre.lower().endswith('.docx'): nombre += ".docx"

        c_contratos, _ = Carpeta.objects.get_or_create(nombre="Contratos Generados", cliente=cliente, padre=None)
        nuevo = Documento(cliente=cliente, carpeta=c_contratos, nombre_archivo=nombre, subido_por=request.user)
        nuevo.archivo.save(nombre, ContentFile(buffer.getvalue()))
        nuevo.save()
        Bitacora.objects.create(usuario=request.user, cliente=cliente, accion='generacion', descripcion=f"Gener√≥ contrato: {nombre}")
        return redirect('visor_docx', documento_id=nuevo.id)

    return render(request, 'generador/llenar.html', {'cliente': cliente, 'plantilla': plantilla, 'variables': formulario})

@login_required
def visor_docx(request, documento_id):
    doc = get_object_or_404(Documento, id=documento_id)
    html = ""
    if doc.nombre_archivo.endswith('.docx'):
        try:
            with doc.archivo.open() as f: html = mammoth.convert_to_html(f).value
        except: pass
    return render(request, 'generador/visor.html', {'doc': doc, 'contenido_html': html})

@login_required
def subir_plantilla(request):
    if request.user.rol == 'admin' and request.method == 'POST':
        Plantilla.objects.create(nombre=request.POST.get('nombre'), archivo=request.FILES.get('archivo'))
    return redirect('dashboard')

@login_required
def dise√±ador_plantillas(request):
    if not request.user.access_disenador: return redirect('dashboard')
    
    if request.method == 'POST':
        nombre = request.POST.get('nombre')
        archivo = request.FILES.get('archivo_base')
        data_reemplazos = request.POST.get('reemplazos')

        if archivo and nombre:
            try:
                doc = DocumentoWord(archivo)
                if data_reemplazos:
                    lista = json.loads(data_reemplazos)
                    for item in lista:
                        original = item.get('texto_original', '')
                        variable = "{{ " + item.get('variable', '') + " }}" 
                        for p in doc.paragraphs:
                            if original in p.text: p.text = p.text.replace(original, variable)
                        for table in doc.tables:
                            for row in table.rows:
                                for cell in row.cells:
                                    for p in cell.paragraphs:
                                        if original in p.text: p.text = p.text.replace(original, variable)

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                nombre_archivo = nombre if nombre.endswith('.docx') else f"{nombre}.docx"
                nueva_plantilla = Plantilla(nombre=nombre)
                nueva_plantilla.archivo.save(nombre_archivo, ContentFile(buffer.getvalue()))
                nueva_plantilla.save()
                messages.success(request, f"¬°Plantilla '{nombre}' guardada!")
            except Exception as e:
                messages.error(request, f"Error: {e}")
            return redirect('dashboard')
    return render(request, 'generador/dise√±ador.html', {'glosario': VariableEstandar.objects.all().order_by('clave')})

@csrf_exempt
@login_required
def previsualizar_word_raw(request):
    if request.method == 'POST' and request.FILES.get('archivo'):
        try:
            f = request.FILES['archivo']
            result = mammoth.convert_to_html(f)
            return JsonResponse({'html': result.value})
        except Exception as e:
            return JsonResponse({'status': 'error', 'msg': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'msg': 'No se envi√≥ archivo'}, status=400)

@csrf_exempt
@login_required
def crear_variable_api(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body)
            clave = data.get('clave')
            descripcion = data.get('descripcion', '')
            tipo = data.get('tipo', 'texto')
            if not clave: return JsonResponse({'status': 'error', 'msg': 'Falta la clave'}, status=400)
            variable, created = VariableEstandar.objects.get_or_create(clave=clave, defaults={'descripcion': descripcion, 'tipo': tipo})
            return JsonResponse({'status': 'ok', 'id': str(variable.id), 'clave': variable.clave, 'created': created})
        except Exception as e:
            return JsonResponse({'status': 'error', 'msg': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'msg': 'M√©todo no permitido'}, status=405)

@csrf_exempt
def api_convertir_html(request):
    import weasyprint 
    if request.method == 'POST':
        try:
            try: data = json.loads(request.body); html_content = data.get('html', '')
            except: html_content = request.POST.get('html', '')
            if not html_content: return JsonResponse({'error': 'No content'}, status=400)
            response = HttpResponse(content_type='application/pdf')
            response['Content-Disposition'] = 'attachment; filename="documento_dise√±ado.pdf"'
            base_url = request.build_absolute_uri('/')
            weasyprint.HTML(string=html_content, base_url=base_url).write_pdf(response)
            return response
        except Exception as e: return JsonResponse({'error': str(e)}, status=500)
    return JsonResponse({'status': 'error', 'message': 'Only POST allowed'}, status=405)

# ==========================================
# 7. COTIZACIONES Y SERVICIOS
# ==========================================

@login_required
def gestion_servicios(request):
    if not request.user.access_cotizaciones: return redirect('dashboard')
    servicios = Servicio.objects.all().order_by('nombre')
    return render(request, 'cotizaciones/servicios.html', {'servicios': servicios})

@login_required
def guardar_servicio(request):
    if request.method == 'POST':
        s_id = request.POST.get('servicio_id')
        s = get_object_or_404(Servicio, id=s_id) if s_id else Servicio()
        
        s.nombre = request.POST.get('nombre')
        s.descripcion = request.POST.get('descripcion')
        s.precio_base = request.POST.get('precio')
        
        # Guardar lista de campos din√°micos (Nombre: Valor)
        nombres = request.POST.getlist('campo_nombre[]')
        valores = request.POST.getlist('campo_valor[]')
        
        estructura = []
        for nombre, valor in zip(nombres, valores):
            if nombre.strip():
                estructura.append({'nombre': nombre.strip(), 'valor': valor.strip()})
        
        s.campos_dinamicos = estructura
        s.save()
        messages.success(request, "Servicio actualizado correctamente.")
    return redirect('gestion_servicios')

@login_required
def eliminar_servicio(request, servicio_id):
    get_object_or_404(Servicio, id=servicio_id).delete()
    return redirect('gestion_servicios')

@login_required
def lista_cotizaciones(request):
    if not request.user.access_cotizaciones: return redirect('dashboard')
    return render(request, 'cotizaciones/lista.html', {'cotizaciones': Cotizacion.objects.all().order_by('-fecha_creacion')})

# En expedientes/views.py
@login_required
def nueva_cotizacion(request):
    if request.method == 'POST':
        # 1. Datos Generales
        titulo = request.POST.get('titulo')
        
        # 2. Datos Cliente
        prospecto_empresa = request.POST.get('prospecto_empresa')
        prospecto_nombre = request.POST.get('prospecto_nombre')
        prospecto_email = request.POST.get('prospecto_email')
        prospecto_telefono = request.POST.get('prospecto_telefono')
        prospecto_direccion = request.POST.get('prospecto_direccion')
        prospecto_cargo = request.POST.get('prospecto_cargo')
        validez = request.POST.get('validez_hasta')
        
        # T√≠tulo autom√°tico si viene vac√≠o
        if not titulo:
            cliente_ref = prospecto_empresa if prospecto_empresa else prospecto_nombre
            titulo = f"Cotizaci√≥n para {cliente_ref}"

        # 3. Descuento
        porcentaje_str = request.POST.get('porcentaje_descuento', '0')
        try:
            porcentaje_descuento = Decimal(porcentaje_str)
        except:
            porcentaje_descuento = Decimal('0.00')

        # 4. L√ìGICA DE IVA FLEXIBLE
        aplica_iva = request.POST.get('aplica_iva') == 'on'
        
        # Capturamos la tasa personalizada del formulario
        # Si no viene o hay error, usamos 16 por defecto
        tasa_str = request.POST.get('porcentaje_iva_personalizado', '16')
        try:
            tasa_iva = Decimal(tasa_str)
        except:
            tasa_iva = Decimal('16.00')

        # 5. Crear Objeto Cotizaci√≥n
        cotizacion = Cotizacion.objects.create(
            titulo=titulo,
            prospecto_empresa=prospecto_empresa,
            prospecto_nombre=prospecto_nombre,
            prospecto_email=prospecto_email,
            prospecto_telefono=prospecto_telefono,
            prospecto_direccion=prospecto_direccion,
            prospecto_cargo=prospecto_cargo,
            porcentaje_descuento=porcentaje_descuento,
            validez_hasta=validez if validez else None,
            
            # Guardamos configuraci√≥n de IVA
            aplica_iva=aplica_iva,
            porcentaje_iva=tasa_iva,  # <--- AQU√ç SE GUARDA LA TASA (8, 16, etc)
            
            creado_por=request.user
        )

        # 6. Procesar Servicios (Items)
        servicios_ids = request.POST.getlist('servicios_seleccionados')
        cantidades = request.POST.getlist('cantidades')
        precios = request.POST.getlist('precios_personalizados')
        descripciones = request.POST.getlist('descripciones_personalizadas')

        for s_id, cant, prec, desc in zip(servicios_ids, cantidades, precios, descripciones):
            if s_id:
                servicio = get_object_or_404(Servicio, id=s_id)
                cantidad = int(cant)
                try:
                    precio_u = Decimal(prec)
                except:
                    precio_u = Decimal('0.00')
                
                ItemCotizacion.objects.create(
                    cotizacion=cotizacion,
                    servicio=servicio,
                    cantidad=cantidad,
                    precio_unitario=precio_u,
                    descripcion_personalizada=desc
                )
        
        # 7. Calcular Totales Finales
        cotizacion.calcular_totales()

        messages.success(request, 'Cotizaci√≥n creada exitosamente.')
        return redirect('detalle_cotizacion', cotizacion_id=cotizacion.id)

    # GET: Mostrar formulario
    servicios = Servicio.objects.all()
    return render(request, 'cotizaciones/crear.html', {'servicios': servicios})

@login_required
def detalle_cotizacion(request, cotizacion_id):
    c = get_object_or_404(Cotizacion, id=cotizacion_id)
    return render(request, 'cotizaciones/detalle.html', {'c': c, 'plantillas_ws': PlantillaMensaje.objects.filter(tipo='whatsapp')})

@login_required
def generar_pdf_cotizacion(request, cotizacion_id):
    import weasyprint
    c = get_object_or_404(Cotizacion, id=cotizacion_id)
    html = render_to_string('cotizaciones/pdf_template.html', {'c': c, 'base_url': request.build_absolute_uri('/')})
    response = HttpResponse(content_type='application/pdf')
    weasyprint.HTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf(response)
    return response

@login_required
def convertir_a_cliente(request, cotizacion_id):
    # Imports necesarios para esta l√≥gica espec√≠fica
    from django.core.files.base import ContentFile
    import weasyprint
    
    c = get_object_or_404(Cotizacion, id=cotizacion_id)
    
    # 1. Validaci√≥n: Si ya es cliente, redirigir
    if c.cliente_convertido:
        messages.warning(request, f"Esta cotizaci√≥n ya pertenece al cliente {c.cliente_convertido}")
        return redirect('detalle_cliente', cliente_id=c.cliente_convertido.id)

    # 2. Buscar o Crear Cliente
    nombre_busqueda = c.prospecto_empresa if c.prospecto_empresa else c.prospecto_nombre
    cli = Cliente.objects.filter(nombre_empresa__iexact=nombre_busqueda).first()

    if not cli:
        # Crear Cliente Nuevo
        cli = Cliente.objects.create(
            nombre_empresa=nombre_busqueda,
            nombre_contacto=c.prospecto_nombre,
            email=c.prospecto_email,
            telefono=c.prospecto_telefono,
            datos_extra={'direccion': c.prospecto_direccion, 'cargo': c.prospecto_cargo}
        )
        # Asignar permisos si no es admin
        if request.user.rol != 'admin':
            request.user.clientes_asignados.add(cli)

    # 3. Buscar la Carpeta "Cotizaciones"
    # Usamos "Cotizaciones" (May√∫scula) para coincidir con el Signal
    carpeta_db, _ = Carpeta.objects.get_or_create(
        nombre="Cotizaciones",
        cliente=cli,
        defaults={'es_expediente': False}
    )

    # 4. Generar el PDF en memoria
    html_string = render_to_string('cotizaciones/pdf_template.html', {'c': c})
    html = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri())
    pdf_content = html.write_pdf()

    # 5. Definir nombre del archivo seguro
    nombre_safe = slugify(c.titulo or f"v1_{c.id}").replace("-", "_")
    nombre_archivo = f"Cotizacion_{c.id}_{nombre_safe}.pdf"

    # 6. GUARDAR EL ARCHIVO (Usando modelo DOCUMENTO)
    # Esto soluciona que no apareciera en el Dashboard y evita errores de ruta
    if not Documento.objects.filter(carpeta=carpeta_db, nombre_archivo=nombre_archivo).exists():
        nuevo_doc = Documento(
            cliente=cli,
            carpeta=carpeta_db,
            nombre_archivo=nombre_archivo,
            subido_por=request.user
        )
        # ContentFile guarda los bytes del PDF directamente en el sistema de archivos
        nuevo_doc.archivo.save(nombre_archivo, ContentFile(pdf_content))
        nuevo_doc.save()

    # 7. Registrar en Finanzas (Cuentas por Cobrar)
    # Seleccionamos el monto correcto dependiendo si la cotizaci√≥n llevaba IVA o no
    monto_final_cobro = c.total_con_iva if c.aplica_iva else c.total

    CuentaPorCobrar.objects.create(
        cliente=cli,
        cotizacion=c,
        concepto=f"Cotizaci√≥n #{c.id} - {c.titulo or 'Proyecto'}",
        monto_total=monto_final_cobro,     # <--- Total real (con o sin IVA)
        saldo_pendiente=monto_final_cobro, # Inicialmente se debe todo
        fecha_vencimiento=c.validez_hasta or timezone.now().date()
    )

    # 8. Actualizar Cotizaci√≥n
    c.estado = 'aceptada'
    c.cliente_convertido = cli
    c.save()

    messages.success(request, "Cliente creado/asociado, PDF guardado y cuenta por cobrar generada.")
    return redirect('detalle_cliente', cliente_id=cli.id)

# FUNCI√ìN DE CORREO ACTUALIZADA (RESEND + NOMBRE INTELIGENTE)
# ----------------------------------------------------
@login_required
def enviar_cotizacion_email(request, cotizacion_id):
    cotizacion = get_object_or_404(Cotizacion, id=cotizacion_id)
    
    if request.method == 'POST':
        asunto = request.POST.get('asunto')
        mensaje_usuario = request.POST.get('mensaje')
        
        # Datos de la Firma Personalizable
        firma_nombre = request.POST.get('firma_nombre', 'Lic. Maribel Aldana Santos')
        firma_cargo = request.POST.get('firma_cargo', 'Gestiones Corpad | Directora General')
        usar_logo_default = request.POST.get('usar_logo_default') == 'on'
        
        # 1. Renderizar el PDF (para adjuntarlo)
        html_string = render_to_string('cotizaciones/pdf_template.html', {'c': cotizacion})
        html = weasyprint.HTML(string=html_string, base_url=request.build_absolute_uri())
        pdf_file = html.write_pdf()

        # 2. Construir el Cuerpo del Correo (HTML)
        # Aqu√≠ incrustamos tu mensaje y la firma editable
        html_content = f"""
        <html>
            <body style="font-family: Arial, sans-serif; color: #333;">
                <div style="padding: 20px;">
                    <p style="white-space: pre-line;">{mensaje_usuario}</p>
                    <br><br>
                    <div style="border-top: 1px solid #ddd; padding-top: 20px; display: flex; align-items: center;">
                        {'<img src="cid:logo_firma" style="width: 50px; height: 50px; border-radius: 50%; margin-right: 15px;">' if usar_logo_default else ''}
                        <div>
                            <strong style="font-size: 14px; color: #2D1B4B; display: block;">{firma_nombre}</strong>
                            <span style="font-size: 12px; color: #666;">{firma_cargo}</span>
                        </div>
                    </div>
                </div>
            </body>
        </html>
        """
        text_content = strip_tags(html_content)

        # 3. Configurar el Email
        email = EmailMultiAlternatives(
            subject=asunto,
            body=text_content,
            from_email=settings.DEFAULT_FROM_EMAIL,
            to=[cotizacion.prospecto_email]
        )
        email.attach_alternative(html_content, "text/html")

        # 4. Adjuntar PDF
        filename = f"Cotizacion_{cotizacion.id}.pdf"
        email.attach(filename, pdf_file, 'application/pdf')

        # 5. Adjuntar Logo como imagen en l√≠nea (CID) si se solicit√≥
        if usar_logo_default:
            # Ruta a tu logo est√°tico (Aseg√∫rate de que la ruta sea correcta en tu PC)
            logo_path = os.path.join(settings.BASE_DIR, 'static', 'img', 'logo.png') 
            if os.path.exists(logo_path):
                with open(logo_path, 'rb') as f:
                    logo_data = f.read()
                    logo = MIMEImage(logo_data)
                    logo.add_header('Content-ID', '<logo_firma>')
                    email.attach(logo)

        # Enviar
        email.send()
        messages.success(request, f'Correo enviado exitosamente a {cotizacion.prospecto_email}')
        
    return redirect('detalle_cotizacion', cotizacion_id=cotizacion_id)
@login_required
def eliminar_cotizacion(request, cotizacion_id):
    if not request.user.access_cotizaciones:
        messages.error(request, "No tienes permiso para realizar esta acci√≥n.")
        return redirect('lista_cotizaciones')
    
    cotizacion = get_object_or_404(Cotizacion, id=cotizacion_id)
    cotizacion_id_ref = cotizacion.id 
    cotizacion.delete()
    
    messages.success(request, f"La cotizaci√≥n #{cotizacion_id_ref} fue eliminada exitosamente.")
    return redirect('lista_cotizaciones')

# ==========================================
# 8. FINANZAS
# ==========================================

@login_required
def panel_finanzas(request):
    if not request.user.access_finanzas: return redirect('dashboard')
    cuentas = CuentaPorCobrar.objects.all().order_by('-fecha_emision')
    return render(request, 'finanzas/panel.html', {'cuentas': cuentas, 'total_por_cobrar': sum(c.saldo_pendiente for c in cuentas), 'total_cobrado': sum(c.monto_pagado for c in cuentas)})

@login_required
def registrar_pago(request):
    if request.method == 'POST':
        Pago.objects.create(
            cuenta_id=request.POST.get('cuenta_id'), monto=Decimal(request.POST.get('monto')),
            metodo=request.POST.get('metodo'), referencia=request.POST.get('referencia'), registrado_por=request.user
        )
    return redirect('panel_finanzas')

@login_required
def recibo_pago_pdf(request, pago_id):
    import weasyprint
    p = get_object_or_404(Pago, id=pago_id)
    html = render_to_string('finanzas/recibo_template.html', {'p': p, 'base_url': request.build_absolute_uri('/')})
    response = HttpResponse(content_type='application/pdf')
    weasyprint.HTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf(response)
    return response

# ==========================================
# 9. AGENDA
# ==========================================

@login_required
def agenda_legal(request):
    if not request.user.access_agenda: return redirect('dashboard')
    hoy = timezone.now()
    proximas = Evento.objects.filter(tipo='audiencia', inicio__gte=hoy, usuario=request.user).order_by('inicio')[:5]
    clientes = Cliente.objects.all() if request.user.rol == 'admin' else request.user.clientes_asignados.all()
    return render(request, 'agenda/calendario.html', {'clientes': clientes, 'proximas_audiencias': proximas})

@login_required
def api_eventos(request):
    if not request.user.access_agenda: return JsonResponse([], safe=False)
    start, end = request.GET.get('start'), request.GET.get('end')
    qs = Evento.objects.filter(inicio__range=[start, end])
    if request.user.rol != 'admin': qs = qs.filter(Q(usuario=request.user) | Q(cliente__in=request.user.clientes_asignados.all()))
    eventos = []
    for e in qs:
        titulo = f"{e.cliente.nombre_empresa}: {e.titulo}" if e.cliente else e.titulo
        eventos.append({'id': e.id, 'title': titulo, 'start': e.inicio.isoformat(), 'end': e.fin.isoformat() if e.fin else None, 'backgroundColor': e.color_hex, 'extendedProps': {'descripcion': e.descripcion, 'tipo': e.get_tipo_display()}})
    return JsonResponse(eventos, safe=False)

@csrf_exempt
@login_required
def mover_evento_api(request):
    if request.method == 'POST':
        try:
            data = json.loads(request.body); evento = get_object_or_404(Evento, id=data.get('id'))
            if request.user.rol != 'admin' and evento.usuario != request.user: return JsonResponse({'status': 'error', 'msg': 'Sin permiso'})
            evento.inicio = data.get('start')
            if data.get('end'): evento.fin = data.get('end')
            evento.save(); return JsonResponse({'status': 'ok'})
        except Exception as e: return JsonResponse({'status': 'error', 'msg': str(e)})
    return JsonResponse({'status': 'error'})

@login_required
def crear_evento(request):
    if request.method == 'POST':
        inicio = timezone.make_aware(timezone.datetime.strptime(f"{request.POST.get('fecha')} {request.POST.get('hora')}", "%Y-%m-%d %H:%M"))
        cliente = get_object_or_404(Cliente, id=request.POST.get('cliente_id')) if request.POST.get('cliente_id') else None
        Evento.objects.create(usuario=request.user, titulo=request.POST.get('titulo'), inicio=inicio, tipo=request.POST.get('tipo'), cliente=cliente, descripcion=request.POST.get('descripcion'))
        messages.success(request, "Evento agendado.")
    return redirect('agenda_legal')

@login_required
def eliminar_evento(request, evento_id):
    evento = get_object_or_404(Evento, id=evento_id)
    if request.user.rol == 'admin' or evento.usuario == request.user:
        evento.delete(); return JsonResponse({'status': 'ok'})
    return JsonResponse({'status': 'error'}, status=403)
# En expedientes/views.py

@login_required
def eliminar_plantilla(request, plantilla_id):
    if request.user.rol != 'admin':
        messages.error(request, "No tienes permisos.")
        return redirect('dashboard')
        
    plantilla = get_object_or_404(Plantilla, id=plantilla_id)
    nombre = plantilla.nombre
    
    # Borrar archivo f√≠sico y registro
    plantilla.archivo.delete() 
    plantilla.delete()
    
    messages.success(request, f"Plantilla '{nombre}' eliminada.")
    
    # Intentar volver a la p√°gina anterior
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))
def subir_archivo_requisito(request, carpeta_id):
    if request.method == 'POST':
        carpeta = get_object_or_404(Carpeta, id=carpeta_id)
        archivo = request.FILES.get('archivo')
        nombre_requisito = request.POST.get('nombre_requisito') # Aqu√≠ recibimos "ACTA CONSTITUTIVA", etc.

        if archivo and nombre_requisito:
            # 1. Borrar si ya exist√≠a uno anterior con ese nombre (para reemplazar)
            Documento.objects.filter(carpeta=carpeta, nombre_archivo=nombre_requisito).delete()

            # 2. Crear el nuevo documento renombrado
            nuevo_doc = Documento(
                cliente=carpeta.cliente,
                carpeta=carpeta,
                archivo=archivo,
                nombre_archivo=nombre_requisito, # ¬°Aqu√≠ ocurre la magia del renombrado!
                subido_por=request.user
            )
            nuevo_doc.save()
            messages.success(request, f'Se carg√≥ correctamente: {nombre_requisito}')
        else:
            messages.error(request, 'Error al subir el archivo.')
            
        return redirect('detalle_cliente', cliente_id=carpeta.cliente.id)
    return redirect('dashboard')
def enviar_recordatorio_documentacion(request, cliente_id):
    cliente = get_object_or_404(Cliente, id=cliente_id)
    
    # 1. Escaneamos qu√© falta (Solo lo que est√° en Rojo)
    faltantes_por_carpeta = {}
    total_faltantes = 0
    
    for carpeta in cliente.carpetas_drive.all():
        detalle = carpeta.obtener_detalle_cumplimiento()
        if detalle:
            # Filtramos solo los que tienen estado 'missing'
            items_rojos = [item['nombre'] for item in detalle if item['estado'] == 'missing']
            if items_rojos:
                faltantes_por_carpeta[carpeta.nombre] = items_rojos
                total_faltantes += len(items_rojos)
    
    # 2. Si no falta nada, avisamos y no enviamos correo
    if total_faltantes == 0:
        messages.success(request, "¬°Este cliente ya tiene toda su documentaci√≥n completa! No es necesario enviar recordatorios.")
        return redirect('detalle_cliente', cliente_id=cliente.id)

    # 3. Redacci√≥n del Correo Formal
    asunto = f"Pendientes de Documentaci√≥n - {cliente.nombre_empresa} - AppLegal"
    
    mensaje = f"""
Estimado(a) {cliente.nombre_contacto},

Esperamos que este correo le encuentre bien.

Le escribimos para darle seguimiento a su expediente de regularizaci√≥n. Para poder avanzar con los tr√°mites ante las autoridades correspondientes, hemos detectado que a√∫n tenemos algunos documentos pendientes de recibir.

A continuaci√≥n, le compartimos el listado de los requisitos faltantes organizados por carpeta:
------------------------------------------------------------
"""

    for nombre_carpeta, documentos in faltantes_por_carpeta.items():
        mensaje += f"\nüìÇ {nombre_carpeta}:\n"
        for doc in documentos:
            mensaje += f"   [ ] {doc}\n"

    mensaje += f"""
------------------------------------------------------------

Le agradecer√≠amos mucho si pudiera compartirnos estos archivos a la brevedad posible, ya sea subi√©ndolos directamente a la plataforma o respondiendo a este correo.

Si tiene alguna duda sobre alg√∫n requisito en espec√≠fico, quedamos totalmente a sus √≥rdenes para apoyarle.

Atentamente,

Gestiones Cordpad
"""

    # 4. Env√≠o del Correo
    try:
        if cliente.email:
            send_mail(
                asunto,
                mensaje,
                settings.DEFAULT_FROM_EMAIL, # Aseg√∫rate de tener esto configurado en settings.py
                [cliente.email],
                fail_silently=False,
            )
            messages.success(request, f"‚úÖ Se envi√≥ el recordatorio a {cliente.email} con {total_faltantes} documentos faltantes.")
        else:
            messages.warning(request, "‚ö†Ô∏è El cliente no tiene un correo electr√≥nico registrado.")
    except Exception as e:
        messages.error(request, f"‚ùå Error al enviar el correo: {str(e)}")

    return redirect('detalle_cliente', cliente_id=cliente.id)
@login_required
def mover_archivo_drive(request, archivo_id):
    doc = get_object_or_404(Documento, id=archivo_id)
    
    # Verificamos permisos
    if not (request.user.can_edit_client or request.user.can_upload_files or request.user.rol == 'admin'):
        messages.error(request, "No tienes permiso para mover archivos.")
        return redirect('detalle_cliente', cliente_id=doc.cliente.id)

    if request.method == 'POST':
        destino_id = request.POST.get('carpeta_destino')
        
        if destino_id == 'ROOT':
            doc.carpeta = None # Mover a Ra√≠z
            nombre_destino = "Carpeta Ra√≠z"
        else:
            carpeta_destino = get_object_or_404(Carpeta, id=destino_id)
            doc.carpeta = carpeta_destino
            nombre_destino = carpeta_destino.nombre
            
        doc.save()
        messages.success(request, f"Archivo movido a: {nombre_destino}")
        
    # Redirigir a donde est√°bamos
    return redirect(request.META.get('HTTP_REFERER', 'dashboard'))
@login_required
def generador_qr(request):
    qr_url = None
    
    # Valores por defecto
    data = ""
    color_fill = "#2D1B4B" # Morado oscuro de tu marca
    color_back = "#FFFFFF" # Blanco

    if request.method == 'POST':
        data = request.POST.get('data')
        color_fill = request.POST.get('color_fill', '#2D1B4B')
        color_back = request.POST.get('color_back', '#FFFFFF')

        if data:
            # Configuraci√≥n del QR
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_H,
                box_size=10,
                border=4,
            )
            qr.add_data(data)
            qr.make(fit=True)

            # Generar imagen con colores personalizados
            img = qr.make_image(
                fill_color=color_fill, 
                back_color=color_back
            )

            # Convertir a base64 para mostrar en HTML sin guardar archivo
            buffer = BytesIO()
            img.save(buffer, format="PNG")
            img_str = base64.b64encode(buffer.getvalue()).decode()
            qr_url = f"data:image/png;base64,{img_str}"

    return render(request, 'generador_qr.html', {
        'qr_url': qr_url,
        'data_input': data,
        'color_fill': color_fill,
        'color_back': color_back
    })
@login_required
def buscar_cliente_api(request):
    query = request.GET.get('q', '')
    if len(query) < 2:
        return JsonResponse([], safe=False)
    
    # Buscamos en cotizaciones anteriores empresas que se parezcan
    # Usamos 'distinct' para no traer repetidos
    resultados = Cotizacion.objects.filter(
        Q(prospecto_empresa__icontains=query) | 
        Q(prospecto_nombre__icontains=query)
    ).values(
        'prospecto_empresa', 
        'prospecto_nombre', 
        'prospecto_email', 
        'prospecto_telefono',
        'prospecto_direccion',
        'prospecto_cargo'
    ).distinct()[:5] # Limitamos a 5 sugerencias

    return JsonResponse(list(resultados), safe=False)
# En expedientes/views.py

@login_required
def generar_orden_cobro(request, cuenta_id, tipo_pago):
    import weasyprint
    from django.utils import timezone
    
    cuenta = get_object_or_404(CuentaPorCobrar, id=cuenta_id)
    cotizacion = cuenta.cotizacion
    
    # 1. Capturar datos bancarios de la URL (GET request)
    datos_bancarios = {
        'banco': request.GET.get('banco', 'BBVA'),
        'cuenta': request.GET.get('cuenta_num', ''),
        'clabe': request.GET.get('clabe', ''),
        'titular': request.GET.get('titular', '')
    }

    # 2. C√°lculos Financieros
    # Nota: cuenta.monto_total ya incluye IVA si la cotizaci√≥n lo ten√≠a.
    total_proyecto = cuenta.monto_total 
    
    if tipo_pago == 'anticipo':
        titulo_doc = "ORDEN DE PAGO - ANTICIPO"
        # El 50% del total (incluyendo impuestos si aplica)
        monto_a_pagar = total_proyecto / Decimal(2)
        nota = "Concepto: 50% de anticipo para inicio de gestiones administrativas."
        porcentaje_pago = 50
    else: # liquidacion
        titulo_doc = "ORDEN DE PAGO - LIQUIDACI√ìN"
        # El saldo restante real
        monto_a_pagar = cuenta.saldo_pendiente 
        nota = "Concepto: Pago final contra entrega de resultados."
        porcentaje_pago = 100 if cuenta.monto_pagado == 0 else 50 # Estimado

    context = {
        'cuenta': cuenta,
        'c': cotizacion, # Pasamos la cotizaci√≥n para ver items e IVA
        'titulo_doc': titulo_doc,
        'monto_a_pagar': monto_a_pagar,
        'nota': nota,
        'tipo_pago': tipo_pago,
        'porcentaje_pago': porcentaje_pago,
        'banco': datos_bancarios,
        'fecha_emision': timezone.now(),
        'base_url': request.build_absolute_uri('/')
    }

    html = render_to_string('finanzas/orden_cobro_pdf.html', context)
    response = HttpResponse(content_type='application/pdf')
    filename = f"Cobro_{tipo_pago}_{cuenta.cliente.nombre_empresa}.pdf"
    response['Content-Disposition'] = f'inline; filename="{filename}"'
    weasyprint.HTML(string=html, base_url=request.build_absolute_uri('/')).write_pdf(response)
    return response